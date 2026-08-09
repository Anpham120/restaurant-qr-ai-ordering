# -*- coding: utf-8 -*-
"""Chuyển mục Tài liệu tham khảo sang IEEE, và đổi trích dẫn trong thân sang [n].

Sửa THẲNG tệp .docx, không xuất lại từ markdown — vì bản .docx đang mang các
chỉnh tay (bảng biểu) không có trong markdown, và xuất lại sẽ xoá mất chúng.

Ba việc:

    1. `(Robertson & Zaragoza, 2009)`  ->  `[1]`      trong thân
    2. Danh sách nguồn viết lại theo IEEE, kèm liên kết
    3. Liên kết đặt thành HYPERLINK thật, bấm được trong Word

Cách sửa run: python-docx không cho thay chữ trên cả đoạn mà giữ định dạng, nên
hàm `_thay_giu_dinh_dang` gộp các run lại rồi ghi đè run đầu và xoá phần còn
lại. Đoạn văn xuôi trong báo cáo dùng một định dạng đồng nhất nên cách này an
toàn; nó KHÔNG an toàn với đoạn có chữ đậm/nghiêng xen giữa, nên chỗ nào cần
giữ thì phải xử riêng.

Chạy:  python sua_tham_khao_ieee.py
"""
from __future__ import annotations

import re
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

HERE = Path(__file__).resolve().parent
TEP = HERE / "output" / "BAO_CAO_HOC_MAY_KPDL.docx"

# Thứ tự [1]..[9] GIỮ NGUYÊN thứ tự đang có trong báo cáo, để không phải sửa
# những chỗ người đọc đã quen. IEEE cho phép đánh số theo thứ tự xuất hiện lần
# đầu; giữ thứ tự cũ là lựa chọn ít xáo trộn nhất.
NGUON = [
    ('S. Robertson and H. Zaragoza, "The probabilistic relevance framework: BM25 and beyond," '
     '*Foundations and Trends in Information Retrieval*, vol. 3, no. 4, pp. 333–389, 2009.',
     "https://doi.org/10.1561/1500000019"),
    ('J. Chen, S. Xiao, P. Zhang, K. Luo, D. Lian, and Z. Liu, "BGE M3-Embedding: Multi-lingual, '
     'multi-functionality, multi-granularity text embeddings through self-knowledge distillation," '
     "*arXiv preprint* arXiv:2402.03216, 2024.",
     "https://arxiv.org/abs/2402.03216"),
    ('L. Wang, N. Yang, X. Huang, B. Jiao, L. Yang, D. Jiang, R. Majumder, and F. Wei, '
     '"Text embeddings by weakly-supervised contrastive pre-training," *arXiv preprint* '
     "arXiv:2212.03533, 2022.",
     "https://arxiv.org/abs/2212.03533"),
    ('G. V. Cormack, C. L. A. Clarke, and S. Buettcher, "Reciprocal rank fusion outperforms '
     'Condorcet and individual rank learning methods," in *Proc. 32nd Int. ACM SIGIR Conf. on '
     "Research and Development in Information Retrieval (SIGIR '09)*, 2009, pp. 758–759.",
     "https://doi.org/10.1145/1571941.1572114"),
    ('P. Lewis, E. Perez, A. Piktus, F. Petroni, V. Karpukhin, N. Goyal, H. Küttler, M. Lewis, '
     'W. Yih, T. Rocktäschel, S. Riedel, and D. Kiela, "Retrieval-augmented generation for '
     'knowledge-intensive NLP tasks," in *Advances in Neural Information Processing Systems '
     "(NeurIPS)*, vol. 33, 2020, pp. 9459–9474.",
     "https://arxiv.org/abs/2005.11401"),
    ('K. Järvelin and J. Kekäläinen, "Cumulated gain-based evaluation of IR techniques," '
     "*ACM Transactions on Information Systems*, vol. 20, no. 4, pp. 422–446, 2002.",
     "https://doi.org/10.1145/582415.582418"),
    ('N. Reimers and I. Gurevych, "Sentence-BERT: Sentence embeddings using Siamese '
     "BERT-networks,\" in *Proc. 2019 Conf. on Empirical Methods in Natural Language Processing "
     "(EMNLP-IJCNLP)*, 2019, pp. 3982–3992.",
     "https://arxiv.org/abs/1908.10084"),
    ('Q. McNemar, "Note on the sampling error of the difference between correlated proportions '
     'or percentages," *Psychometrika*, vol. 12, no. 2, pp. 153–157, 1947.',
     "https://doi.org/10.1007/BF02295996"),
    ('E. B. Wilson, "Probable inference, the law of succession, and statistical inference," '
     "*Journal of the American Statistical Association*, vol. 22, no. 158, pp. 209–212, 1927.",
     "https://doi.org/10.2307/2276774"),
]

# `(Tác giả, năm)` -> `[n]`. Khớp cả dạng có "et al." lẫn dạng hai tác giả.
DOI_CHIEU = [
    (re.compile(r"\(Robertson\s*&\s*Zaragoza,?\s*2009\)"), "[1]"),
    (re.compile(r"\(Chen\s+et al\.?,?\s*2024\)"), "[2]"),
    (re.compile(r"\(Wang\s+et al\.?,?\s*2022\)"), "[3]"),
    (re.compile(r"\(Cormack\s+et al\.?,?\s*2009\)"), "[4]"),
    (re.compile(r"\(Lewis\s+et al\.?,?\s*2020\)"), "[5]"),
    (re.compile(r"\(Järvelin\s*&\s*Kekäläinen,?\s*2002\)"), "[6]"),
    (re.compile(r"\(Reimers\s*&\s*Gurevych,?\s*2019\)"), "[7]"),
    (re.compile(r"\(McNemar,?\s*1947\)"), "[8]"),
    (re.compile(r"\(Wilson,?\s*1927\)"), "[9]"),
]

# Chỗ nêu tên một khái niệm mà CHƯA có trích dẫn nào — thêm [n] ngay sau cụm.
# Chỉ thêm ở LẦN XUẤT HIỆN ĐẦU của mỗi nguồn, đúng quy ước IEEE.
THEM_CHIEU = [
    ("Reciprocal Rank Fusion", "[4]"),
    ("nDCG@k = DCG@k", "[6]"),
    ("kiểm định McNemar ghép cặp", "[8]"),
    ("phương pháp Wilson", "[9]"),
    ("sentence-transformers", "[7]"),
]


def _thay_giu_dinh_dang(p, cu: str, moi: str) -> bool:
    """Thay chuỗi trong một đoạn, giữ định dạng của run ĐẦU TIÊN."""
    if cu not in p.text:
        return False
    runs = p.runs
    if not runs:
        return False
    moi_text = p.text.replace(cu, moi, 1)
    runs[0].text = moi_text
    for r in runs[1:]:
        r.text = ""
    return True


def _lien_ket(p, url: str, chu: str | None = None) -> None:
    """Chèn một siêu liên kết bấm được vào cuối đoạn."""
    phan = p.part
    r_id = phan.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    h = docx.oxml.shared.OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    r = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")
    mau = docx.oxml.shared.OxmlElement("w:color")
    mau.set(qn("w:val"), "0563C1")
    gach = docx.oxml.shared.OxmlElement("w:u")
    gach.set(qn("w:val"), "single")
    fonts = docx.oxml.shared.OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    sz = docx.oxml.shared.OxmlElement("w:sz")
    sz.set(qn("w:val"), "26")                       # 13pt = 26 half-points
    for x in (fonts, mau, gach, sz):
        rPr.append(x)
    r.append(rPr)
    t = docx.oxml.shared.OxmlElement("w:t")
    t.text = chu or url
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    h.append(r)
    p._p.append(h)


def _dat_chu(p, chu: str, dam: bool = False, nghieng: bool = False) -> None:
    r = p.add_run(chu)
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    r.bold = dam
    r.italic = nghieng
    r.font.color.rgb = RGBColor(0, 0, 0)


def _viet_muc(p, van: str) -> None:
    """Ghi một mục nguồn, hiểu `*nghiêng*` cho tên tạp chí."""
    for k, phan in enumerate(van.split("*")):
        if phan:
            _dat_chu(p, phan, nghieng=(k % 2 == 1))


def main() -> int:
    d = docx.Document(str(TEP))

    # --- 1. Trích dẫn trong thân -------------------------------------------------
    doi = 0
    for p in d.paragraphs:
        for rx, so in DOI_CHIEU:
            if rx.search(p.text):
                _thay_giu_dinh_dang(p, rx.search(p.text).group(0), so)
                doi += 1
    for p in d.paragraphs:
        for cum, so in THEM_CHIEU:
            if cum in p.text and so not in p.text:
                _thay_giu_dinh_dang(p, cum, f"{cum} {so}")
                doi += 1
                break
    print(f"  trích dẫn trong thân: {doi} chỗ -> dạng [n]")

    # --- 2. Danh sách nguồn ------------------------------------------------------
    i = next(k for k, p in enumerate(d.paragraphs) if "THAM KHẢO" in p.text.upper())
    muc = [p for p in d.paragraphs[i + 1:i + 1 + len(NGUON)]]
    assert len(muc) == len(NGUON), f"tìm thấy {len(muc)} mục, cần {len(NGUON)}"

    for k, (p, (van, url)) in enumerate(zip(muc, NGUON), 1):
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        p.style = d.styles["Normal"]
        p.paragraph_format.left_indent = Pt(28)
        p.paragraph_format.first_line_indent = Pt(-28)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _dat_chu(p, f"[{k}]\t")
        _viet_muc(p, van)
        _dat_chu(p, " ")
        _lien_ket(p, url)
    print(f"  danh sách nguồn: {len(NGUON)} mục -> IEEE, mỗi mục một liên kết")

    d.save(str(TEP))
    print(f"đã lưu {TEP.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
