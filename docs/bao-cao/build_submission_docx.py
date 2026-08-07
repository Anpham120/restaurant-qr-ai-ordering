from __future__ import annotations

import html as html_lib
import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import html as lxml_html


ROOT = Path(__file__).resolve().parents[2]
SOURCE = Path(__file__).with_name("BAO_CAO_CONG_NGHE_PHAN_MEM.md")
OUTPUT_DIR = Path(__file__).with_name("output")
OUTPUT = OUTPUT_DIR / "BAO_CAO_CONG_NGHE_PHAN_MEM.docx"
DIAGRAM_DIR = OUTPUT_DIR / "_generated_diagrams"
REPOSITORY_WEB_ROOT = "https://github.com/Anpham120/restaurant-qr-ai-ordering"
REPOSITORY_WEB_BRANCH = "develop"

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_LEFT_CM = 3.0
MARGIN_RIGHT_CM = 2.0
MARGIN_TOP_CM = 2.0
MARGIN_BOTTOM_CM = 2.0
CONTENT_WIDTH_IN = (PAGE_WIDTH_CM - MARGIN_LEFT_CM - MARGIN_RIGHT_CM) / 2.54
CONTENT_WIDTH_DXA = round(CONTENT_WIDTH_IN * 1440)

INK = "111111"
MUTED = "606060"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
CALLOUT = "EEF3F7"
BORDER = "A6A6A6"

TABLE_CAPTIONS = [
    "Danh sách thành viên nhóm",
    "Thống kê tổng quan dự án",
    "Danh mục từ viết tắt",
    "Phân công công việc của các thành viên",
    "Mối liên hệ giữa đề tài và nội dung học phần",
    "Các thành phần của Product Vision",
    "Kết quả mong đợi và tiêu chí hoàn thành MVP",
    "Phạm vi chức năng của MVP",
    "Phương pháp thu thập yêu cầu",
    "Persona khách hàng Minh",
    "Persona nhân viên quầy Lan",
    "Persona bếp trưởng Tuấn",
    "User stories, acceptance criteria và minh chứng",
    "Product Backlog của MVP",
    "Yêu cầu chức năng",
    "Yêu cầu phi chức năng",
    "Trách nhiệm và giới hạn của các tầng kiến trúc",
    "So sánh modular monolith và microservices",
    "Phân tách trách nhiệm giữa backend nghiệp vụ và dịch vụ AI",
    "Bất biến dữ liệu và cơ chế cưỡng chế",
    "Trạng thái phiên bàn và hành vi khi quét lại",
    "Đánh giá ưu, nhược điểm của công nghệ được lựa chọn",
    "So sánh kết quả các phương pháp truy hồi",
    "Cấu trúc kho tri thức",
    "So sánh fine-tune với RAG kết hợp lọc nhãn",
    "Các tầng kiểm thử và phạm vi kết luận",
    "Ma trận truy vết yêu cầu – kiểm thử",
    "Các lớp bảo mật và bằng chứng",
    "Quy tắc code review và quản lý mã nguồn",
    "Definition of Done",
    "Kế hoạch và kết quả các milestone",
    "Hệ thống nhãn của GitHub Issues",
    "Thống kê commit và pull request",
    "Các phiên bản phát hành",
    "Cấu hình branch ruleset",
    "Nhật ký sử dụng công cụ AI",
    "Đóng góp của từng thành viên",
    "Tổng hợp liên hệ lý thuyết Sommerville",
    "Quy mô mã nguồn theo thành phần",
    "Kết quả đo chất lượng",
    "Ma trận rubric và minh chứng",
    "Mức độ đáp ứng mục tiêu đề tài",
    "Hạn chế của sản phẩm và quá trình phát triển",
]

FIGURE_CAPTIONS = [
    "Kiến trúc tổng thể của hệ thống",
    "Mô hình quan hệ dữ liệu rút gọn",
    "Máy trạng thái của phiên bàn",
    "Dòng chảy CI/CD",
    "Các milestone của repository",
    "Danh sách issue kèm nhãn và người được gán",
    "Hoạt động commit theo tuần",
    "Danh sách pull request đã hợp nhất",
    "Lịch sử chạy GitHub Actions",
    "Các phiên bản phát hành trên GitHub",
    "Branch ruleset của main và develop",
    "Nhóm giao diện khách hàng và vận hành",
    "Điểm vào gọi món trên thiết bị di động",
    "So sánh gợi ý AI trước và sau khi khách nêu dị ứng",
    "Bảng bếp thời gian thực",
    "Quầy thu ngân và hóa đơn phiên bàn",
]


def font_path(bold: bool = False, italic: bool = False) -> str:
    fonts = Path(r"C:\Windows\Fonts")
    if bold and italic:
        return str(fonts / "arialbi.ttf")
    if bold:
        return str(fonts / "arialbd.ttf")
    if italic:
        return str(fonts / "ariali.ttf")
    return str(fonts / "arial.ttf")


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(font_path(bold=bold), size=size)


def clean_text(text: str) -> str:
    text = html_lib.unescape(text)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    text = text.replace("`", "")
    text = text.replace("→", "→")
    return "".join(ch for ch in text if ch == "\n" or ch == "\t" or ord(ch) >= 32).strip()


def clean_inline_markup(text: str) -> str:
    """Remove HTML wrappers while preserving Markdown links and emphasis."""
    text = html_lib.unescape(text)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    return "".join(
        ch for ch in text if ch == "\n" or ch == "\t" or ord(ch) >= 32
    ).strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def estimate_widths(rows: list[list[str]], total: int) -> list[int]:
    cols = max(len(row) for row in rows)
    max_lens = []
    for col in range(cols):
        lengths = [len(clean_text(row[col])) if col < len(row) else 0 for row in rows]
        max_lens.append(max(8, min(max(lengths, default=8), 70)))
    weights = [math.sqrt(length) for length in max_lens]
    minimum = 620 if cols >= 5 else 760
    available = total - minimum * cols
    if available < 0:
        minimum = max(420, total // cols // 2)
        available = total - minimum * cols
    weight_sum = sum(weights)
    widths = [minimum + round(available * weight / weight_sum) for weight in weights]
    widths[-1] += total - sum(widths)
    return widths


def set_run_font(run, name="Times New Roman", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text: str, url: str, bold=False, italic=False, code=False):
    part = paragraph.part
    relationship = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    r_fonts = OxmlElement("w:rFonts")
    chosen = "Consolas" if code else "Times New Roman"
    r_fonts.set(qn("w:ascii"), chosen)
    r_fonts.set(qn("w:hAnsi"), chosen)
    r_fonts.set(qn("w:eastAsia"), chosen)
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_TOKEN = re.compile(
    r"(\[((?:\[[^\]]+\])|[^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|`([^`]+)`|\*([^*]+)\*)"
)


def resolve_hyperlink(url: str) -> str | None:
    if url.startswith(("http://", "https://")):
        return url
    if not url or url.startswith("#"):
        return None

    relative_url, separator, fragment = url.partition("#")
    target = (SOURCE.parent / relative_url).resolve()
    try:
        repository_path = target.relative_to(ROOT).as_posix()
    except ValueError:
        return None

    link_kind = "tree" if target.is_dir() else "blob"
    resolved = (
        f"{REPOSITORY_WEB_ROOT}/{link_kind}/{REPOSITORY_WEB_BRANCH}/{repository_path}"
    )
    if separator and fragment:
        resolved = f"{resolved}#{fragment}"
    return resolved


def add_inline(paragraph, text: str, size=None, base_bold=False, base_italic=False) -> None:
    pos = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, bold=base_bold, italic=base_italic)
        token = match.group(0)
        if token.startswith("["):
            label, url = match.group(2), match.group(3)
            resolved_url = resolve_hyperlink(url)
            if resolved_url:
                add_hyperlink(paragraph, clean_text(label), resolved_url)
            else:
                run = paragraph.add_run(clean_text(label))
                set_run_font(run, size=size, bold=base_bold, italic=base_italic)
        elif token.startswith("**"):
            run = paragraph.add_run(match.group(4))
            set_run_font(run, size=size, bold=True, italic=base_italic)
        elif token.startswith("`"):
            run = paragraph.add_run(match.group(5))
            set_run_font(run, name="Consolas", size=(size or 10), bold=base_bold, italic=base_italic)
            run.font.color.rgb = RGBColor.from_string("7A1F1F")
        else:
            run = paragraph.add_run(match.group(6))
            set_run_font(run, size=size, bold=base_bold, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, bold=base_bold, italic=base_italic)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Cập nhật mục lục trong Microsoft Word: Ctrl+A, sau đó nhấn F9."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])


def add_field(paragraph, instruction_text: str, placeholder_text: str = ""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = instruction_text
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = placeholder_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])
    return run


def add_front_matter_heading(doc: Document, text: str, include_in_toc: bool = False) -> None:
    p = doc.add_paragraph(style="Heading 2" if include_in_toc else None)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text.upper())
    set_run_font(run, size=16, bold=True)


def bookmark_name(label: str, number: int) -> str:
    prefix = "bang" if label == "Bảng" else "hinh"
    return f"{prefix}_{number}"


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_list_of_captions(doc: Document, label: str, titles: list[str]) -> None:
    for number, title in enumerate(titles, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(CONTENT_WIDTH_IN),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        entry = p.add_run(f"{label} {number} — {title}\t")
        set_run_font(entry, size=10)
        page = add_field(p, f"PAGEREF {bookmark_name(label, number)} \\h", "0")
        set_run_font(page, size=10)


def add_seq_caption(doc: Document, label: str, title: str, number_value: int) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = False
    prefix = p.add_run(f"{label} ")
    set_run_font(prefix, size=10, italic=True)
    number = add_field(p, f"SEQ {label} \\* ARABIC", "0")
    set_run_font(number, size=10, italic=True)
    suffix = p.add_run(f" — {title}")
    set_run_font(suffix, size=10, italic=True)
    bookmark_id = (1000 if label == "Bảng" else 2000) + number_value
    add_bookmark(p, bookmark_name(label, number_value), bookmark_id)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, clean_inline_markup(text), size=10, base_italic=True)
    p.paragraph_format.keep_with_next = False


def resolve_image(src: str) -> Path:
    src = html_lib.unescape(src)
    return (SOURCE.parent / src).resolve()


def add_picture(doc: Document, path: Path, max_width=CONTENT_WIDTH_IN, max_height=7.5, caption=None):
    if not path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Không tìm thấy hình: {path.name}]")
        set_run_font(run, italic=True, color="9B1C1C")
        return
    with Image.open(path) as img:
        width_px, height_px = img.size
    ratio = width_px / max(height_px, 1)
    width = min(max_width, max_height * ratio)
    height = width / max(ratio, 0.01)
    if height > max_height:
        height = max_height
        width = height * ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
    if caption:
        add_caption(doc, caption)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=cols)
    table.style = "Table Grid"
    table.alignment = 0
    table.autofit = False
    widths = estimate_widths(normalized, CONTENT_WIDTH_DXA)
    apply_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    font_size = 9.5 if cols <= 3 else 8.5 if cols == 4 else 7.5

    for r_idx, row in enumerate(normalized):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(
                p,
                clean_inline_markup(value),
                size=font_size,
                base_bold=(r_idx == 0),
            )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for idx, line in enumerate(lines):
        if idx == 1 and re.match(r"^\|\s*:?-+", line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_html_gallery(doc: Document, block: str) -> None:
    root = lxml_html.fragment_fromstring(block, create_parent="div")
    html_rows = root.xpath(".//tr")
    if not html_rows:
        for image in root.xpath(".//img"):
            src = image.get("src")
            if src:
                add_picture(doc, resolve_image(src))
        return
    col_count = max(len(row.xpath("./td")) for row in html_rows)
    table = doc.add_table(rows=len(html_rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = False
    widths = [CONTENT_WIDTH_DXA // col_count] * col_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    apply_table_geometry(table, widths)
    for r_idx, row in enumerate(html_rows):
        prevent_row_split(table.rows[r_idx])
        cells = row.xpath("./td")
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if c_idx >= len(cells):
                continue
            node = cells[c_idx]
            image_nodes = node.xpath(".//img")
            if image_nodes:
                src = image_nodes[0].get("src")
                image_path = resolve_image(src)
                if image_path.exists():
                    with Image.open(image_path) as img:
                        ratio = img.width / max(img.height, 1)
                    width = min((CONTENT_WIDTH_IN / col_count) - 0.25, 3.0)
                    height = min(3.25, width / max(ratio, 0.01))
                    if height >= 3.25:
                        width = height * ratio
                    p.add_run().add_picture(str(image_path), width=Inches(width), height=Inches(height))
            text_parts = []
            for element in node.xpath(".//strong|.//sub"):
                value = " ".join(element.itertext()).strip()
                if value:
                    text_parts.append(value)
            if not text_parts:
                value = " ".join(node.itertext()).strip()
                if value:
                    text_parts.append(value)
            for idx, value in enumerate(text_parts):
                value = re.sub(r"^Hình\s+\d+(?:\.\d+)?\s*[—-]\s*", "", value)
                target = p if idx == 0 and not image_nodes else cell.add_paragraph()
                target.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(
                    target,
                    clean_inline_markup(value),
                    size=8.5,
                    base_bold=(idx == 0),
                )
    doc.add_paragraph()


def add_code_block(doc: Document, code: str, language: str = "") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    apply_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    for idx, line in enumerate(code.splitlines() or [""]):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.5)
    doc.add_paragraph()


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    apply_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, clean_text(text), size=11)
    doc.add_paragraph()


def draw_box(draw, xy, text, fill="#EEF3F7", outline="#5B7C99", bold=False, font_size=26):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=3)
    font = pil_font(font_size, bold=bold)
    lines = text.split("\n")
    line_heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(line_heights) + 8 * (len(lines) - 1)
    y = (y1 + y2 - total_h) / 2
    for line, height in zip(lines, line_heights):
        box = draw.textbbox((0, 0), line, font=font)
        width = box[2] - box[0]
        draw.text(((x1 + x2 - width) / 2, y), line, font=font, fill="#111111")
        y += height + 8


def arrow(draw, start, end, color="#395B75", width=5):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 14
    left = (
        end[0] - size * math.cos(angle - math.pi / 6),
        end[1] - size * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - size * math.cos(angle + math.pi / 6),
        end[1] - size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=color)


def save_diagram_1(path: Path):
    img = Image.new("RGB", (1800, 1500), "white")
    d = ImageDraw.Draw(img)
    title = pil_font(34, bold=True)
    d.text((900, 35), "KIẾN TRÚC TỔNG THỂ", font=title, fill="#1F4E79", anchor="ma")
    clients = [
        (80, 120, 530, 270, "customer-web\ngiới thiệu · thực đơn"),
        (675, 120, 1125, 270, "ordering-web\ngọi món · chat AI"),
        (1270, 120, 1720, 270, "admin-web\nadmin · quầy · bếp"),
    ]
    for x1, y1, x2, y2, label in clients:
        draw_box(d, (x1, y1, x2, y2), label, fill="#E7F1F8", font_size=27)
        arrow(d, ((x1 + x2) / 2, y2), (900, 390))
    draw_box(d, (570, 390, 1230, 540), "ASP.NET Core\nREST API · JWT/RBAC · SignalR", fill="#DDEBF7", bold=True)
    draw_box(d, (100, 680, 650, 830), "PostgreSQL 16\n24 bảng · 22 migration", fill="#E2F0D9", bold=True)
    draw_box(d, (760, 680, 1310, 830), "Dịch vụ AI · FastAPI\n/v1/chat · /v1/chat/stream", fill="#FFF2CC", bold=True)
    draw_box(d, (1350, 680, 1720, 830), "LLM Gateway\n9router", fill="#FCE4D6")
    arrow(d, (750, 540), (430, 680))
    arrow(d, (1050, 540), (1035, 680))
    arrow(d, (1310, 755), (1350, 755))
    draw_box(d, (680, 1010, 1070, 1145), "Hiểu câu hỏi\n→ ràng buộc nhãn", fill="#F8F2E4")
    draw_box(d, (170, 1250, 670, 1390), "Lọc thực đơn theo nhãn\nmã tất định", fill="#E2F0D9", bold=True)
    draw_box(d, (720, 1250, 1160, 1390), "Truy hồi embedding\nkho tri thức", fill="#E7E6F7")
    draw_box(d, (1210, 1250, 1690, 1390), "Viết câu trả lời\n+ hàng rào món/giá", fill="#FCE4D6")
    arrow(d, (1035, 830), (875, 1010))
    arrow(d, (800, 1145), (420, 1250))
    arrow(d, (930, 1145), (940, 1250))
    arrow(d, (1160, 1320), (1210, 1320))
    img.save(path)


def save_diagram_2(path: Path):
    img = Image.new("RGB", (1800, 1250), "white")
    d = ImageDraw.Draw(img)
    d.text((900, 35), "MÔ HÌNH QUAN HỆ DỮ LIỆU RÚT GỌN", font=pil_font(34, True), fill="#1F4E79", anchor="ma")
    columns = [
        ["RestaurantTable", "TableSession", "Order", "OrderItem", "OrderStatusHistory"],
        ["TableSessionCartItem", "TableInvoice", "Payment", "PaymentTransaction", "CounterShift"],
        ["Category", "MenuItem", "MenuItemKnowledge", "ChatSession", "ChatMessage"],
    ]
    xs = [90, 650, 1210]
    boxes = {}
    for col_idx, names in enumerate(columns):
        for row_idx, name in enumerate(names):
            x1, y1 = xs[col_idx], 140 + row_idx * 205
            box = (x1, y1, x1 + 470, y1 + 115)
            boxes[name] = box
            draw_box(d, box, name, fill=["#E7F1F8", "#E2F0D9", "#FFF2CC"][col_idx], font_size=25)
    pairs = [
        ("RestaurantTable", "TableSession"),
        ("TableSession", "Order"),
        ("Order", "OrderItem"),
        ("Order", "OrderStatusHistory"),
        ("TableSession", "TableSessionCartItem"),
        ("TableSession", "TableInvoice"),
        ("TableInvoice", "Payment"),
        ("Payment", "PaymentTransaction"),
        ("Category", "MenuItem"),
        ("MenuItem", "MenuItemKnowledge"),
        ("MenuItem", "OrderItem"),
        ("ChatSession", "ChatMessage"),
    ]
    for source, target in pairs:
        s, t = boxes[source], boxes[target]
        start = ((s[0] + s[2]) / 2, s[3]) if abs(s[0] - t[0]) < 50 else (s[2], (s[1] + s[3]) / 2)
        end = ((t[0] + t[2]) / 2, t[1]) if abs(s[0] - t[0]) < 50 else (t[0], (t[1] + t[3]) / 2)
        arrow(d, start, end, width=3)
    img.save(path)


def save_diagram_3(path: Path):
    img = Image.new("RGB", (1800, 1500), "white")
    d = ImageDraw.Draw(img)
    d.text((900, 35), "MÁY TRẠNG THÁI PHIÊN BÀN", font=pil_font(34, True), fill="#1F4E79", anchor="ma")
    states = [
        ("New", "Quét QR / mở phiên"),
        ("CartPending", "Có món trong giỏ"),
        ("OrderInProgress", "Đã gửi bếp / có thể gọi thêm"),
        ("ReadyForPayment", "Mọi món đã phục vụ"),
        ("PaymentPending", "Đã yêu cầu thanh toán"),
        ("Paid", "Quầy xác nhận đã thu"),
    ]
    boxes = []
    for idx, (state, note) in enumerate(states):
        x1, y1 = 470, 120 + idx * 215
        box = (x1, y1, 1330, y1 + 125)
        boxes.append(box)
        draw_box(d, box, f"{state}\n{note}", fill="#E7F1F8" if idx < 3 else "#E2F0D9", bold=True, font_size=25)
        if idx:
            arrow(d, ((boxes[idx - 1][0] + boxes[idx - 1][2]) / 2, boxes[idx - 1][3]), ((x1 + 1330) / 2, y1))
    d.text((1380, 385), "gọi thêm lượt", font=pil_font(22), fill="#606060")
    d.arc((1220, 310, 1660, 650), 260, 100, fill="#395B75", width=4)
    d.text((120, 1040), "Hủy yêu cầu thanh toán", font=pil_font(22), fill="#606060")
    arrow(d, (470, 1135), (250, 1135), width=4)
    arrow(d, (250, 1135), (470, 920), width=4)
    d.text((900, 1430), "Paid → đóng phiên → bàn trở về trạng thái trống", font=pil_font(26, True), fill="#1F4E79", anchor="ma")
    img.save(path)


def save_diagram_4(path: Path):
    img = Image.new("RGB", (2000, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((1000, 40), "DÒNG CHẢY CI/CD", font=pil_font(34, True), fill="#1F4E79", anchor="ma")
    labels = [
        "Pull Request\n→ develop",
        "CI · 5 job\nbuild · test · E2E",
        "Security\nCodeQL · Trivy",
        "Deploy\nstaging",
        "Promote\n→ main",
        "Deploy\nproduction",
    ]
    boxes = []
    x = 40
    for idx, label in enumerate(labels):
        width = 285 if idx != 1 else 360
        box = (x, 250, x + width, 430)
        boxes.append(box)
        draw_box(d, box, label, fill="#E7F1F8" if idx < 3 else "#E2F0D9", bold=True, font_size=24)
        if idx:
            arrow(d, (boxes[idx - 1][2], 340), (box[0], 340))
        x += width + 45
    rollback = (1500, 610, 1910, 770)
    draw_box(d, rollback, "rollback.yml\nkhi smoke test thất bại", fill="#FCE4D6", bold=True, font_size=24)
    arrow(d, ((boxes[-1][0] + boxes[-1][2]) / 2, boxes[-1][3]), ((rollback[0] + rollback[2]) / 2, rollback[1]))
    img.save(path)


def generate_diagrams() -> list[Path]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    paths = [DIAGRAM_DIR / f"diagram-{idx}.png" for idx in range(1, 5)]
    save_diagram_1(paths[0])
    save_diagram_2(paths[1])
    save_diagram_3(paths[2])
    save_diagram_4(paths[3])
    return paths


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.3

    heading_specs = {
        "Heading 1": (16, 16, 8, True),
        "Heading 2": (14, 12, 6, False),
        "Heading 3": (13, 10, 4, False),
        "Heading 4": (12, 8, 3, False),
    }
    for name, (size, before, after, page_break) in heading_specs.items():
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = page_break

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(12)
        style.paragraph_format.left_indent = Cm(0.75)
        style.paragraph_format.first_line_indent = Cm(-0.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    # Keep the three-level table of contents compact enough to avoid an
    # almost-empty trailing page while retaining every heading.
    for level, size in ((1, 9.5), (2, 9.0), (3, 8.5)):
        name = f"toc {level}"
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style._element.set(qn("w:styleId"), f"TOC{level}")
        style.base_style = styles["Normal"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.left_indent = Cm(0.42 * (level - 1))


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(MARGIN_TOP_CM)
        section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
        section.left_margin = Cm(MARGIN_LEFT_CM)
        section.right_margin = Cm(MARGIN_RIGHT_CM)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)
        section.different_first_page_header_footer = True

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("INFO2005 · CMC RESTAURANT")
        set_run_font(run, size=9, color=MUTED)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Trang ")
        set_run_font(run, size=9, color=MUTED)
        add_page_field(p)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    logo = ROOT / "frontend" / "src" / "mocks" / "images" / "logo.png"
    if logo.exists():
        p.add_run().add_picture(str(logo), width=Inches(1.25))

    for text, size, bold, after in [
        ("TRƯỜNG ĐẠI HỌC CMC", 14, True, 2),
        ("KHOA CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG", 13, True, 26),
        ("BÁO CÁO BÀI TẬP LỚN", 22, True, 8),
        ("HỌC PHẦN CÔNG NGHỆ PHẦN MỀM", 16, True, 20),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    add_inline(
        p,
        "Đề tài: CMC Restaurant — Hệ thống gọi món và theo dõi trạng thái đơn tại bàn bằng QR, tích hợp trợ lý AI tư vấn thực đơn",
        size=16,
        base_bold=True,
    )

    metadata = [
        ("Mã học phần", "INFO2005"),
        ("Số tín chỉ", "03"),
        ("Giảng viên phụ trách", "Trương Anh Hoàng"),
        ("Nhóm thực hiện", "05 sinh viên"),
        ("Thời gian thực hiện", "04/06/2026 – 02/08/2026"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.autofit = False
    table.style = "Table Grid"
    apply_table_geometry(table, [2600, CONTENT_WIDTH_DXA - 2600])
    for idx, (label, value) in enumerate(metadata):
        set_cell_shading(table.cell(idx, 0), LIGHT_BLUE)
        for col, text in enumerate((label, value)):
            p = table.cell(idx, col).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, text, size=11, base_bold=(col == 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run("Hà Nội, năm 2026")
    set_run_font(run, size=12, italic=True)
    doc.add_page_break()


def is_structure(line: str, next_line: str = "") -> bool:
    stripped = line.strip()
    return (
        not stripped
        or stripped.startswith("#")
        or stripped.startswith("```")
        or stripped.startswith(">")
        or stripped == "---"
        or stripped.startswith("<table")
        or stripped.startswith("<div")
        or stripped.startswith("<img")
        or re.match(r"^[-*]\s+", stripped) is not None
        or re.match(r"^\d+\.\s+", stripped) is not None
        or (stripped.startswith("|") and next_line.strip().startswith("|"))
        or re.match(r"^!\[[^\]]*\]\([^)]+\)", stripped) is not None
    )


def build() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = generate_diagrams()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("**Nhóm thực hiện"))
    i = start
    paragraph_buffer: list[str] = []
    mermaid_idx = 0
    table_idx = 0
    figure_idx = 0
    skip_toc = False
    skip_manual_list = False

    def flush_paragraph():
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        if text:
            p = doc.add_paragraph()
            add_inline(p, text, size=12)
        paragraph_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        next_line = lines[i + 1] if i + 1 < len(lines) else ""

        if stripped == "## Mục lục":
            flush_paragraph()
            add_front_matter_heading(doc, "MỤC LỤC")
            toc = doc.add_paragraph()
            add_toc_field(toc)
            skip_toc = True
            i += 1
            continue
        if skip_toc:
            if stripped in ("## Danh mục bảng", "## Danh mục hình", "## Danh mục từ viết tắt"):
                skip_toc = False
            else:
                i += 1
                continue

        if stripped == "## Danh mục bảng":
            flush_paragraph()
            add_front_matter_heading(doc, "DANH MỤC BẢNG", include_in_toc=True)
            add_list_of_captions(doc, "Bảng", TABLE_CAPTIONS)
            skip_manual_list = True
            i += 1
            continue

        if stripped == "## Danh mục hình":
            flush_paragraph()
            add_front_matter_heading(doc, "DANH MỤC HÌNH", include_in_toc=True)
            add_list_of_captions(doc, "Hình", FIGURE_CAPTIONS)
            skip_manual_list = True
            i += 1
            continue

        if skip_manual_list:
            if stripped.startswith("## "):
                skip_manual_list = False
            else:
                i += 1
                continue

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped in ("---", "</div>") or stripped.startswith("<div"):
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("<sub>") and stripped.endswith("</sub>"):
            flush_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, clean_inline_markup(stripped), size=9, base_italic=True)
            i += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            language = stripped[3:].strip()
            block_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block_lines.append(lines[i])
                i += 1
            if language == "mermaid":
                if figure_idx >= len(FIGURE_CAPTIONS):
                    raise ValueError("Thiếu tiêu đề trong FIGURE_CAPTIONS.")
                add_picture(
                    doc,
                    diagrams[mermaid_idx],
                    max_width=CONTENT_WIDTH_IN,
                    max_height=7.2,
                )
                add_seq_caption(doc, "Hình", FIGURE_CAPTIONS[figure_idx], figure_idx + 1)
                mermaid_idx += 1
                figure_idx += 1
            else:
                add_code_block(doc, "\n".join(block_lines), language)
            i += 1
            continue

        if stripped.startswith("<table"):
            flush_paragraph()
            block = [line]
            i += 1
            while i < len(lines):
                block.append(lines[i])
                if "</table>" in lines[i]:
                    break
                i += 1
            gallery_html = "\n".join(block)
            if len(re.findall(r"<tr\b", gallery_html, flags=re.I)) > 1:
                doc.add_page_break()
            add_html_gallery(doc, gallery_html)
            if figure_idx >= len(FIGURE_CAPTIONS):
                raise ValueError("Thiếu tiêu đề trong FIGURE_CAPTIONS.")
            add_seq_caption(doc, "Hình", FIGURE_CAPTIONS[figure_idx], figure_idx + 1)
            figure_idx += 1
            i += 1
            continue

        html_image = re.search(r'<img\s+[^>]*src="([^"]+)"[^>]*>', stripped)
        if html_image:
            flush_paragraph()
            add_picture(doc, resolve_image(html_image.group(1)), max_width=3.3, max_height=6.0)
            if figure_idx >= len(FIGURE_CAPTIONS):
                raise ValueError("Thiếu tiêu đề trong FIGURE_CAPTIONS.")
            add_seq_caption(doc, "Hình", FIGURE_CAPTIONS[figure_idx], figure_idx + 1)
            figure_idx += 1
            i += 1
            continue

        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            flush_paragraph()
            add_picture(doc, resolve_image(image_match.group(2)), caption=None)
            if figure_idx >= len(FIGURE_CAPTIONS):
                raise ValueError("Thiếu tiêu đề trong FIGURE_CAPTIONS.")
            add_seq_caption(doc, "Hình", FIGURE_CAPTIONS[figure_idx], figure_idx + 1)
            figure_idx += 1
            i += 1
            continue

        if stripped.startswith("|") and next_line.strip().startswith("|") and re.search(r"-{3,}", next_line):
            flush_paragraph()
            table_lines = [line, next_line]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if table_idx >= len(TABLE_CAPTIONS):
                raise ValueError("Thiếu tiêu đề trong TABLE_CAPTIONS.")
            add_seq_caption(doc, "Bảng", TABLE_CAPTIONS[table_idx], table_idx + 1)
            add_markdown_table(doc, parse_markdown_table(table_lines))
            table_idx += 1
            continue

        if stripped.startswith("*Hình"):
            flush_paragraph()
            while i < len(lines):
                current = lines[i].strip()
                i += 1
                if current.endswith("*"):
                    break
            continue

        if stripped.startswith("*Ghi chú về minh chứng đóng góp"):
            flush_paragraph()
            note_lines = []
            while i < len(lines):
                current = lines[i].strip()
                note_lines.append(current)
                i += 1
                if current.endswith("*"):
                    break
            p = doc.add_paragraph()
            add_inline(
                p,
                clean_text(" ".join(note_lines).strip("*")),
                size=9.5,
                base_italic=True,
            )
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(re.sub(r"^>\s?", "", lines[i].strip()))
                i += 1
            add_callout(doc, " ".join(quote_lines))
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = clean_text(heading.group(2))
            style_level = min(4, level)
            p = doc.add_paragraph(text, style=f"Heading {style_level}")
            if text in ("Tóm tắt dự án", "Danh mục từ viết tắt", "Bảng phân công công việc"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if text in ("Danh mục từ viết tắt", "Bảng phân công công việc"):
                p.paragraph_format.page_break_before = True
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        number = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if bullet or number:
            flush_paragraph()
            content = bullet.group(1) if bullet else number.group(2)
            continuation = []
            j = i + 1
            while j < len(lines):
                candidate = lines[j]
                if not candidate.strip():
                    break
                if is_structure(candidate, lines[j + 1] if j + 1 < len(lines) else ""):
                    break
                if candidate.startswith("  ") or candidate.startswith("\t"):
                    continuation.append(candidate.strip())
                    j += 1
                else:
                    break
            if continuation:
                content += " " + " ".join(continuation)
            if bullet:
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, content, size=12)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.28)
                p.paragraph_format.first_line_indent = Inches(-0.22)
                add_inline(p, f"{number.group(1)}. {content}", size=12)
            i = j
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            flush_paragraph()
            add_caption(doc, stripped.strip("*"))
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()

    if table_idx != len(TABLE_CAPTIONS):
        raise ValueError(f"Số bảng không khớp: tạo {table_idx}, khai báo {len(TABLE_CAPTIONS)}.")
    if figure_idx != len(FIGURE_CAPTIONS):
        raise ValueError(f"Số hình không khớp: tạo {figure_idx}, khai báo {len(FIGURE_CAPTIONS)}.")

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    core = doc.core_properties
    core.title = "Báo cáo Công nghệ phần mềm — CMC Restaurant"
    core.subject = "INFO2005 — Bài tập lớn"
    core.author = "Nhóm CMC Restaurant"
    core.keywords = "công nghệ phần mềm, QR ordering, AI, RAG, CI/CD"

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    result = build()
    print(result)
