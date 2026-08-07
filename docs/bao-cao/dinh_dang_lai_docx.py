"""Định dạng lại một tệp .docx có sẵn mà KHÔNG đụng tới nội dung chữ.

Dùng cho bản báo cáo đã được sửa tay trong Word: script chỉ đặt lại kiểu chữ
đề mục, màu, căn lề trong bảng, bề rộng cột, chiều cao hàng và xoá đầu/chân
trang. Không thêm, không bớt, không sửa một ký tự nào của văn bản.

    python docs/bao-cao/dinh_dang_lai_docx.py "output/BAO_CAO_PHAN_MEM_HOAN_CHINH.docx"
"""

from __future__ import annotations

import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "Times New Roman"
DEN = RGBColor(0x00, 0x00, 0x00)
VIEN = "BFBFBF"
NEN_TIEU_DE = "EDEDED"

RONG_NOI_DUNG_CM = 16.0
CAO_HANG_TIEU_DE_CM = 0.9
CAO_HANG_CM = 0.75

# cỡ, đậm, nghiêng, IN HOA, căn giữa, thụt lề (cm)
# Mọi đề mục chỉ dùng in đậm — không nghiêng. Phân cấp thể hiện bằng số thứ tự
# và mức thụt lề.
KIEU_DE_MUC = {
    "Heading 1": (14, True, False, True, True, 0.0),
    "Heading 2": (13, True, False, False, False, 0.0),
    "Heading 3": (13, True, False, False, False, 0.5),
    "Heading 4": (13, True, False, False, False, 0.5),
}


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), str(v))
    return e


def _bo_con(cha, ten):
    for x in cha.findall(qn(ten)):
        cha.remove(x)


# --------------------------------------------------------------- đề mục

def dat_de_muc(doc: Document) -> int:
    for ten, (co, dam, ng, hoa, giua, thut) in KIEU_DE_MUC.items():
        if ten not in [s.name for s in doc.styles]:
            continue
        s = doc.styles[ten]
        s.font.name = FONT
        s.font.size = Pt(co)
        s.font.bold = dam
        s.font.italic = ng
        s.font.color.rgb = DEN
        rpr = s.element.get_or_add_rPr()
        _bo_con(rpr, "w:caps")
        if hoa:
            rpr.append(_el("w:caps", **{"w:val": "true"}))
        pf = s.paragraph_format
        pf.alignment = (WD_ALIGN_PARAGRAPH.CENTER if giua
                        else WD_ALIGN_PARAGRAPH.LEFT)
        pf.line_spacing = 1.3
        pf.left_indent = Cm(thut)
        pf.keep_with_next = True

    # Định dạng trực tiếp trên run sẽ đè lên style, nên phải ép lại từng run.
    n = 0
    for p in doc.paragraphs:
        if p.style.name not in KIEU_DE_MUC:
            continue
        co, dam, ng, hoa, giua, thut = KIEU_DE_MUC[p.style.name]
        p.paragraph_format.alignment = (WD_ALIGN_PARAGRAPH.CENTER if giua
                                        else WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.left_indent = Cm(thut)
        for r in p.runs:
            r.font.name = FONT
            r.font.size = Pt(co)
            r.bold = dam
            r.italic = ng
            r.font.color.rgb = DEN
        n += 1
    return n


# --------------------------------------------------------------- bảng

def _chu(x: str) -> str:
    return re.sub(r"\s+", " ", x).strip()


def rong_cot(bang, co_chu: float) -> list[float]:
    n = len(bang.columns)
    rong_ky_tu = 0.0176 * co_chu
    dem = 0.36
    tran = RONG_NOI_DUNG_CM * (0.46 if n <= 3 else 0.34)

    toi_thieu, mong_muon = [], []
    for i in range(n):
        o = [_chu(r.cells[i].text) for r in bang.rows if i < len(r.cells)]
        nhan = o[0] if o else ""
        tu = [w for x in o for w in re.split(r"[\s/,;()]+", x) if w]
        dai = max((len(w) for w in tu), default=1)
        rong_tu = dai * rong_ky_tu + dem
        rong_nhan = (len(nhan) / 2 + 1.5) * rong_ky_tu + dem
        toi_thieu.append(min(max(rong_tu, rong_nhan), tran))
        tb = sum(len(x) for x in o) / max(len(o), 1)
        mong_muon.append(tb * rong_ky_tu * 0.62 + dem)

    tong = sum(toi_thieu)
    if tong >= RONG_NOI_DUNG_CM:
        return [x * RONG_NOI_DUNG_CM / tong for x in toi_thieu]
    du = RONG_NOI_DUNG_CM - tong
    them = [max(mong_muon[i] - toi_thieu[i], 0) for i in range(n)]
    tt = sum(them)
    if tt <= 0:
        return [x + du / n for x in toi_thieu]
    return [toi_thieu[i] + du * them[i] / tt for i in range(n)]


def dat_bang(doc: Document) -> tuple[int, int]:
    so_bang = so_o = 0
    for bang in doc.tables:
        n = len(bang.columns)
        co = 11 if n <= 4 else (10 if n == 5 else 9)
        w = rong_cot(bang, co)

        pr = bang._tbl.tblPr
        _bo_con(pr, "w:tblBorders")
        b = _el("w:tblBorders")
        for canh in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b.append(_el(f"w:{canh}", **{"w:val": "single", "w:sz": "4",
                                         "w:space": "0", "w:color": VIEN}))
        pr.append(b)
        bang.autofit = False

        for i, c in enumerate(bang.columns):
            c.width = Cm(w[i])

        for j, hang in enumerate(bang.rows):
            trpr = hang._tr.get_or_add_trPr()
            _bo_con(trpr, "w:trHeight")
            _bo_con(trpr, "w:cantSplit")
            trpr.append(_el("w:cantSplit"))
            cao = CAO_HANG_TIEU_DE_CM if j == 0 else CAO_HANG_CM
            trpr.append(_el("w:trHeight", **{"w:val": str(int(cao * 567)),
                                             "w:hRule": "atLeast"}))
            if j == 0:
                _bo_con(trpr, "w:tblHeader")
                trpr.append(_el("w:tblHeader", **{"w:val": "true"}))

            for i, o in enumerate(hang.cells):
                if i < n:
                    o.width = Cm(w[i])
                o.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if j == 0:
                    tcpr = o._tc.get_or_add_tcPr()
                    _bo_con(tcpr, "w:shd")
                    tcpr.append(_el("w:shd", **{"w:val": "clear",
                                                "w:color": "auto",
                                                "w:fill": NEN_TIEU_DE}))
                for p in o.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(2)
                    for r in p.runs:
                        if r.font.size is None:
                            r.font.size = Pt(co)
                        if j == 0:
                            r.bold = True
                    so_o += 1
        so_bang += 1
    return so_bang, so_o


# --------------------------------------------------------------- trang

def xoa_dau_chan(doc: Document) -> None:
    for s in doc.sections:
        for phan in (s.header, s.footer):
            phan.is_linked_to_previous = False
            for p in list(phan.paragraphs):
                for r in list(p.runs):
                    r._r.getparent().remove(r._r)
                p.text = ""


def main(duong: Path) -> None:
    if not duong.exists():
        raise SystemExit(f"Không thấy tệp: {duong}")

    sao_luu = duong.with_name(
        f"{duong.stem}_saoluu_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(duong, sao_luu)
    print(f"Đã sao lưu bản gốc: {sao_luu.name}")

    doc = Document(str(duong))
    truoc_doan = len([p for p in doc.paragraphs if p.text.strip()])
    truoc_chu = sum(len(p.text.split()) for p in doc.paragraphs)

    n_de_muc = dat_de_muc(doc)
    n_bang, n_o = dat_bang(doc)
    xoa_dau_chan(doc)

    sau_doan = len([p for p in doc.paragraphs if p.text.strip()])
    sau_chu = sum(len(p.text.split()) for p in doc.paragraphs)

    doc.save(str(duong))

    print(f"Đề mục đã đặt lại : {n_de_muc}")
    print(f"Bảng đã căn chỉnh : {n_bang} bảng · {n_o} ô căn giữa")
    print(f"Đầu/chân trang    : đã xoá")
    print(f"Nội dung           : {truoc_doan} → {sau_doan} đoạn, "
          f"{truoc_chu} → {sau_chu} từ "
          f"{'(KHÔNG đổi)' if (truoc_doan, truoc_chu) == (sau_doan, sau_chu) else '(CÓ ĐỔI — kiểm tra lại)'}")


if __name__ == "__main__":
    p = Path(sys.argv[1]) if len(sys.argv) > 1 else (
        Path(__file__).resolve().parent / "output" / "BAO_CAO_PHAN_MEM_HOAN_CHINH.docx")
    main(p)
